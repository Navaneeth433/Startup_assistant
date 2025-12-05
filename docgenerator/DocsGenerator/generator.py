import os
from docx import Document
from fpdf import FPDF
import ollama
from .utils import add_heading_paragraph, add_bullets, pdf_add_text, pdf_add_title

# -------------------
# Template-based Document Generator
# -------------------

# NDA Generator
def generate_nda(user, startup, other_party, purpose):
    prompt = f"""
    You are a legal assistant. Generate a simple NDA between {startup.startup_name} (Founder: {user.full_name}) 
    and {other_party} for {purpose}.
    Use clear headings and bullet points for clauses.
    """
    response = ollama.chat(
        model="mistral",
        messages=[{"role": "system", "content": "You are a legal assistant."},
                  {"role": "user", "content": prompt}]
    )
    nda_text = response['message']['content']

    doc = Document()
    add_heading_paragraph(doc, f"NDA Agreement: {startup.startup_name} & {other_party}", "")
    add_heading_paragraph(doc, "Clauses:", nda_text)
    
    file_name = f"NDA_{startup.startup_name}_{other_party}.docx"
    doc.save(file_name)
    return file_name

# MoU Generator
def generate_mou(user, startup, partner_name, purpose):
    prompt = f"""
    You are a legal assistant. Generate a MoU between {startup.startup_name} (Founder: {user.full_name})
    and {partner_name} for {purpose}. Include structured headings and bullet points.
    """
    response = ollama.chat(
        model="mistral",
        messages=[{"role": "system", "content": "You are a legal assistant."},
                  {"role": "user", "content": prompt}]
    )
    mou_text = response['message']['content']

    doc = Document()
    add_heading_paragraph(doc, f"MoU: {startup.startup_name} & {partner_name}", "")
    add_heading_paragraph(doc, "Clauses:", mou_text)
    
    file_name = f"MoU_{startup.startup_name}_{partner_name}.docx"
    doc.save(file_name)
    return file_name

# RTI Draft Generator
def generate_rti(user, startup, authority, subject, purpose):
    prompt = f"""
    Draft an RTI application from {user.full_name} ({startup.startup_name}) to {authority} 
    for subject: {subject}. Include purpose: {purpose}. Format professionally.
    """
    response = ollama.chat(
        model="mistral",
        messages=[{"role": "system", "content": "You are a legal assistant."},
                  {"role": "user", "content": prompt}]
    )
    rti_text = response['message']['content']

    doc = Document()
    add_heading_paragraph(doc, f"RTI Application: {startup.startup_name}", "")
    add_heading_paragraph(doc, "Content:", rti_text)
    
    file_name = f"RTI_{startup.startup_name}_{authority}.docx"
    doc.save(file_name)
    return file_name

# Pitch Deck Generator (PDF with formatting)
def generate_pitch_deck(user, startup):
    prompt = f"""
    Generate a concise pitch deck for {startup.startup_name}.
    Sections: Problem, Solution, Market, Business Model, Team, Vision. Use bullet points.
    """
    response = ollama.chat(
        model="mistral",
        messages=[{"role": "system", "content": "You are a startup assistant."},
                  {"role": "user", "content": prompt}]
    )
    pitch_text = response['message']['content']

    pdf = FPDF()
    pdf.add_page()
    pdf_add_title(pdf, f"Pitch Deck: {startup.startup_name}")
    pdf_add_text(pdf, pitch_text)

    file_name = f"PitchDeck_{startup.startup_name}.pdf"
    pdf.output(file_name)
    return file_name
