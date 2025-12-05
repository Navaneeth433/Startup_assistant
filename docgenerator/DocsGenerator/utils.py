from docx import Document
from fpdf import FPDF
import os

# Word helper: add heading + paragraph
def add_heading_paragraph(doc, heading, text, level=1):
    doc.add_heading(heading, level=level)
    doc.add_paragraph(text)

# Word helper: add bullet points
def add_bullets(doc, bullets):
    for b in bullets:
        doc.add_paragraph(f"• {b}")

# PDF helper: add multi-line text
def pdf_add_text(pdf, text, font="Arial", size=12):
    pdf.set_font(font, size=size)
    pdf.multi_cell(0, 8, text)

# PDF helper: add title
def pdf_add_title(pdf, title, font="Arial", size=16):
    pdf.set_font(font, "B", size)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.ln(5)
